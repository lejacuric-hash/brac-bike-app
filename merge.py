from __future__ import annotations

import json
import re
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document 


BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "places to visit.docx"
GEOJSON_PATH = BASE_DIR / "places to visit.geojson"
OUTPUT_PATH = BASE_DIR / "src" / "final_places.json"


IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".jfif")
IMAGE_LINE_PATTERN = re.compile(
	r"(?:^|\b)(?:image|img|photo|slika)\s*[:\-]\s*(?P<value>.+)$",
	re.IGNORECASE,
)
TOWN_HEADER_PATTERN = re.compile(r"^[^a-z0-9]*(?:[A-ZČĆŽŠĐ\s]+)$")


@dataclass
class StoryEntry:
	name: str
	story: str
	image: str = ""


def normalize_text(value: str) -> str:
	"""Normalize text for reliable cross-language matching."""
	if not value:
		return ""

	value = unicodedata.normalize("NFKD", value)
	value = "".join(ch for ch in value if not unicodedata.combining(ch))
	value = value.lower()
	value = value.replace("&", " and ")
	value = re.sub(r"[^a-z0-9]+", " ", value)
	return re.sub(r"\s+", " ", value).strip()


def strip_leading_numbering(value: str) -> str:
	return re.sub(r"^\s*\d+[\.)\-:]\s*", "", value).strip()


def clean_story_text(value: str) -> str:
	value = value.replace("\xa0", " ")
	return re.sub(r"\n{3,}", "\n\n", value).strip()


def looks_like_image_reference(value: str) -> bool:
	lowered = value.strip().lower()
	return lowered.endswith(IMAGE_EXTENSIONS)


def extract_image_reference(value: str) -> str:
	"""Extract image filename from paragraphs like 'Image: foo.jpg' or plain filename lines."""
	value = value.strip()
	if not value:
		return ""

	match = IMAGE_LINE_PATTERN.search(value)
	if match:
		candidate = match.group("value").strip().strip('"\'')
		return candidate

	if looks_like_image_reference(value):
		return value.strip('"\'')

	return ""


def build_known_name_index(geojson_data: dict) -> dict[str, str]:
	"""Map normalized names to display names using all available POI naming fields."""
	known: dict[str, str] = {}

	for feature in geojson_data.get("features", []):
		properties = feature.get("properties", {})
		for key in ("name", "name-en", "id"):
			value = properties.get(key)
			if not isinstance(value, str):
				continue

			cleaned = value.strip()
			if not cleaned:
				continue

			normalized = normalize_text(cleaned)
			if not normalized:
				continue

			# Keep the longest human-readable variant for a normalized key.
			existing = known.get(normalized, "")
			if len(cleaned) > len(existing):
				known[normalized] = cleaned

	return known


def parse_docx_table_entries(document: Document) -> list[StoryEntry]:
	entries: list[StoryEntry] = []

	for table in document.tables:
		for row in table.rows:
			cells = [c.text.strip() for c in row.cells]
			if not cells or all(not c for c in cells):
				continue

			name = strip_leading_numbering(cells[0])
			story = "\n\n".join(part for part in cells[1:] if part).strip() if len(cells) > 1 else ""
			image = ""

			if len(cells) > 2:
				maybe_image = extract_image_reference(cells[2])
				if maybe_image:
					image = maybe_image

			if not image:
				for cell in cells[1:]:
					maybe_image = extract_image_reference(cell)
					if maybe_image:
						image = maybe_image
						break

			if not name:
				continue

			entries.append(StoryEntry(name=name, story=clean_story_text(story), image=image))

	return entries


def parse_docx_prefix_entries(document: Document, known_name_index: dict[str, str]) -> list[StoryEntry]:
	"""Parse one-line entries where paragraphs begin with POI names followed by descriptions."""
	entries: list[StoryEntry] = []

	known_keys = sorted(known_name_index.keys(), key=len, reverse=True)

	for paragraph in document.paragraphs:
		text = paragraph.text.strip()
		if not text:
			continue

		# Skip section headers like "SUPETAR" or "🏘️ BOL".
		if len(text) <= 40 and TOWN_HEADER_PATTERN.match(text):
			continue

		normalized_paragraph = normalize_text(text)
		if not normalized_paragraph:
			continue

		best_key = ""
		for name_key in known_keys:
			if normalized_paragraph == name_key or normalized_paragraph.startswith(name_key + " "):
				best_key = name_key
				break

		if not best_key:
			continue

		name = known_name_index[best_key]
		image = extract_image_reference(text)
		entries.append(StoryEntry(name=name, story=clean_story_text(text), image=image))

	return entries


def parse_docx_heading_entries(document: Document) -> list[StoryEntry]:
	entries: list[StoryEntry] = []

	current_name = ""
	current_lines: list[str] = []
	current_image = ""

	def flush() -> None:
		nonlocal current_name, current_lines, current_image
		if current_name:
			entries.append(
				StoryEntry(
					name=strip_leading_numbering(current_name),
					story=clean_story_text("\n".join(current_lines)),
					image=current_image,
				)
			)
		current_name = ""
		current_lines = []
		current_image = ""

	for paragraph in document.paragraphs:
		text = paragraph.text.strip()
		if not text:
			if current_lines and current_lines[-1] != "":
				current_lines.append("")
			continue

		style_name = paragraph.style.name if paragraph.style else ""
		is_heading = style_name.lower().startswith("heading")

		if is_heading:
			flush()
			current_name = text
			continue

		maybe_image = extract_image_reference(text)
		if maybe_image:
			current_image = maybe_image
			continue

		if current_name:
			current_lines.append(text)

	flush()
	return entries


def parse_docx_block_entries(document: Document) -> list[StoryEntry]:
	"""Fallback parser for simple docs where each block starts with a title line."""
	lines = [p.text.strip() for p in document.paragraphs]

	entries: list[StoryEntry] = []
	block: list[str] = []

	def flush_block(block_lines: list[str]) -> None:
		if not block_lines:
			return

		name = strip_leading_numbering(block_lines[0])
		content = block_lines[1:]
		image = ""

		cleaned_content: list[str] = []
		for line in content:
			maybe_image = extract_image_reference(line)
			if maybe_image and not image:
				image = maybe_image
				continue
			cleaned_content.append(line)

		story = clean_story_text("\n".join(cleaned_content))
		if name:
			entries.append(StoryEntry(name=name, story=story, image=image))

	for line in lines:
		if line:
			block.append(line)
		else:
			flush_block(block)
			block = []
	flush_block(block)

	return entries


def load_story_entries(docx_path: Path, known_name_index: dict[str, str]) -> list[StoryEntry]:
	document = Document(docx_path)

	table_entries = parse_docx_table_entries(document)
	prefix_entries = parse_docx_prefix_entries(document, known_name_index)
	heading_entries = parse_docx_heading_entries(document)
	block_entries = parse_docx_block_entries(document)

	# Keep the richest set while deduplicating by normalized name.
	candidates = table_entries + prefix_entries + heading_entries + block_entries
	dedup: dict[str, StoryEntry] = {}

	for entry in candidates:
		key = normalize_text(entry.name)
		if not key:
			continue

		previous = dedup.get(key)
		if previous is None:
			dedup[key] = entry
			continue

		# Prefer the entry with the longer story, and preserve image if present.
		better = entry if len(entry.story) > len(previous.story) else previous
		if not better.image:
			better.image = previous.image or entry.image
		dedup[key] = better

	return list(dedup.values())


def build_image_index(images_dir: Path) -> dict[str, str]:
	"""Map normalized filename stems to concrete image filenames."""
	index: dict[str, str] = {}
	if not images_dir.exists():
		return index

	for image_file in images_dir.iterdir():
		if not image_file.is_file() or image_file.suffix.lower() not in IMAGE_EXTENSIONS:
			continue

		stem_key = normalize_text(image_file.stem)
		if stem_key and stem_key not in index:
			index[stem_key] = image_file.name

	return index


def choose_image_for_feature(properties: dict, image_index: dict[str, str], story_image: str) -> str:
	if story_image:
		return story_image

	for candidate in feature_name_candidates(properties):
		key = normalize_text(candidate)
		if key in image_index:
			return image_index[key]

	best_name = ""
	best_ratio = 0.0
	for candidate in feature_name_candidates(properties):
		norm_candidate = normalize_text(candidate)
		if not norm_candidate:
			continue

		for image_key, image_name in image_index.items():
			ratio = SequenceMatcher(None, norm_candidate, image_key).ratio()
			if ratio > best_ratio:
				best_ratio = ratio
				best_name = image_name

	return best_name if best_ratio >= 0.92 else ""


def feature_name_candidates(properties: dict) -> list[str]:
	candidates: list[str] = []

	for key in ("name", "name-en", "id"):
		value = properties.get(key)
		if isinstance(value, str) and value.strip():
			candidates.append(value.strip())

	return candidates


def choose_story_for_feature(
	properties: dict,
	stories_by_name: dict[str, StoryEntry],
	story_keys: list[str],
) -> StoryEntry | None:
	feature_candidates = feature_name_candidates(properties)

	# 1) Exact normalized matching against any known name field.
	for candidate in feature_candidates:
		key = normalize_text(candidate)
		if key in stories_by_name:
			return stories_by_name[key]

	# 2) Fuzzy fallback matching for minor punctuation/wording differences.
	best_key = ""
	best_ratio = 0.0

	for candidate in feature_candidates:
		norm_candidate = normalize_text(candidate)
		if not norm_candidate:
			continue

		for story_key in story_keys:
			ratio = SequenceMatcher(None, norm_candidate, story_key).ratio()
			if ratio > best_ratio:
				best_ratio = ratio
				best_key = story_key

	if best_key and best_ratio >= 0.86:
		return stories_by_name[best_key]

	return None


def enrich_geojson_with_stories(geojson_data: dict, stories: list[StoryEntry], image_index: dict[str, str]) -> dict:
	stories_by_name = {normalize_text(item.name): item for item in stories if normalize_text(item.name)}
	story_keys = list(stories_by_name.keys())

	for feature in geojson_data.get("features", []):
		properties = feature.setdefault("properties", {})
		match = choose_story_for_feature(properties, stories_by_name, story_keys)

		if match:
			properties["story"] = match.story
			properties["image"] = choose_image_for_feature(properties, image_index, match.image)
		else:
			properties["story"] = ""
			properties["image"] = choose_image_for_feature(properties, image_index, "")

	return geojson_data


def main() -> None:
	if not DOCX_PATH.exists():
		raise FileNotFoundError(f"Missing DOCX file: {DOCX_PATH}")
	if not GEOJSON_PATH.exists():
		raise FileNotFoundError(f"Missing GeoJSON file: {GEOJSON_PATH}")

	with GEOJSON_PATH.open("r", encoding="utf-8") as geojson_file:
		geojson_data = json.load(geojson_file)

	known_name_index = build_known_name_index(geojson_data)
	stories = load_story_entries(DOCX_PATH, known_name_index)
	image_index = build_image_index(BASE_DIR / "public" / "images")

	enriched = enrich_geojson_with_stories(geojson_data, stories, image_index)

	OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
	with OUTPUT_PATH.open("w", encoding="utf-8") as output_file:
		json.dump(enriched, output_file, ensure_ascii=False, indent=2)

	print(f"Wrote {OUTPUT_PATH}")
	print(f"Loaded {len(stories)} stories")
	print(f"Updated {len(enriched.get('features', []))} features")


if __name__ == "__main__":
	main()
